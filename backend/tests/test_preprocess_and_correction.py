from io import BytesIO

from docx import Document
from pypdf import PdfWriter

from app.application.services.document_input_preprocessor import DocumentInputPreprocessor
from app.application.services.text_correction_service import TextCorrectionService


def test_preprocessor_passthrough_for_image() -> None:
    preprocessor = DocumentInputPreprocessor()
    image_bytes = b"dummy-image"
    rendered, supplemental = preprocessor.preprocess(image_bytes, "sample.png", "image/png")
    assert rendered == image_bytes
    assert supplemental == ""


def test_text_correction_fills_missing_fields() -> None:
    service = TextCorrectionService()
    normalized = {"line_items": [], "total": None, "tax": None}
    corrected = service.apply(
        normalized,
        "invoice no: INV-2026-01 invoice date: 2026/04/29 currency: jpy vendor code: VEND-001 total: 1200 tax: 100",
    )
    assert corrected["invoice_number"] == "INV-2026-01"
    assert corrected["invoice_date"] == "2026-04-29"
    assert corrected["currency"] == "JPY"
    assert corrected["vendor_code"] == "VEND-001"
    assert corrected["total"] == 1200.0
    assert corrected["tax"] == 100.0


def test_preprocessor_extracts_docx_text() -> None:
    preprocessor = DocumentInputPreprocessor()
    doc = Document()
    doc.add_paragraph("invoice no: INV-DOCX-1")
    doc.add_paragraph("total: 900")
    buf = BytesIO()
    doc.save(buf)
    rendered, supplemental = preprocessor.preprocess(
        buf.getvalue(),
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert rendered
    assert "INV-DOCX-1" in supplemental


def test_preprocessor_raises_value_error_for_invalid_pdf() -> None:
    preprocessor = DocumentInputPreprocessor()
    try:
        preprocessor.preprocess(b"not-a-pdf", "broken.pdf", "application/pdf")
        assert False, "ValueError が発生するべき"
    except ValueError as exc:
        assert "PDF" in str(exc)


def test_preprocessor_rasterizes_pdf() -> None:
    preprocessor = DocumentInputPreprocessor()
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    buf = BytesIO()
    writer.write(buf)

    rendered, supplemental = preprocessor.preprocess(buf.getvalue(), "sample.pdf", "application/pdf")
    assert isinstance(rendered, bytes)
    assert len(rendered) > 0
    assert isinstance(supplemental, str)
